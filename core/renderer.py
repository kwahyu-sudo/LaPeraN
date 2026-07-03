from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy

from .models import LaporanPerdin

_WML_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def _strip_bookmarks(element):
    """Remove w:bookmarkStart/w:bookmarkEnd from a cloned XML element.

    deepcopy duplicates bookmark IDs → invalid OOXML → Word repair dialog.
    Call on every clone before inserting into the document tree.
    """
    for bm in element.iterfind('.//{%s}bookmarkStart' % _WML_NS):
        bm.getparent().remove(bm)
    for bm in element.iterfind('.//{%s}bookmarkEnd' % _WML_NS):
        bm.getparent().remove(bm)


def _replace_text(run, old, new):
    if old in run.text:
        run.text = run.text.replace(old, new)
        return True
    return False


def _replace_in_para(para, replacements: dict):
    for run in para.runs:
        for k, v in replacements.items():
            _replace_text(run, k, str(v))


def _replace_in_doc(doc, replacements: dict):
    for para in doc.paragraphs:
        _replace_in_para(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_para(para, replacements)


def _rebuild_dari_section(doc, pelaksana: list, kepada: str, tembusan: str, hari_tanggal: str):
    """Rebuild P[3] and P[4] (Kepada/Dari/Tembusan/Hari) with dynamic pelaksana count."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    if len(doc.paragraphs) < 5:
        return

    p3, p4 = doc.paragraphs[3], doc.paragraphs[4]
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    from docx.shared import Inches
    for p in (p3, p4):
        p.paragraph_format.left_indent = Inches(0.63)
        p.paragraph_format.right_indent = Inches(0.57)
        p.paragraph_format.first_line_indent = 0

    lines = []
    for i, p in enumerate(pelaksana):
        prefix = "Dari : " if i == 0 else "       "  # 7 spasi rata dengan "Dari : "
        lines.append(f"{prefix}{i+1}. {p.nama}")
    dari_text = "\n".join(lines)
    rest = ""
    if tembusan:
        rest += f"\nTembusan : {tembusan}"
    rest += f"\nHari, Tanggal : {hari_tanggal}"

    if p3.runs:
        p3.runs[0].text = f"Kepada : {kepada}"
        for r in p3.runs[1:]:
            r.text = ""
    if p4.runs:
        p4.runs[0].text = dari_text + rest
        for r in p4.runs[1:]:
            r.text = ""

    # Bottom border under P[4] (garis pemisah header)
    from docx.oxml import OxmlElement
    pPr = p4._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p4._element.insert(0, pPr)
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = pBdr.find(qn('w:bottom'))
    if bottom is None:
        bottom = OxmlElement('w:bottom')
        pBdr.append(bottom)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')


def _rebuild_hasil(doc, hasil: list):
    """Clone or trim hasil paragraphs to match count.

    IMPORTANT: This must run BEFORE _replace_in_doc replaces {{HASIL_N}}
    placeholders, because it detects slots by searching for the literal
    {{HASIL_N}} pattern in paragraph text.
    """
    import re
    # Find paragraphs with {{HASIL_N}} placeholders
    slots = []
    for pi, para in enumerate(doc.paragraphs):
        if re.search(r"\{\{HASIL_\d+\}\}", para.text):
            slots.append(pi)

    if not slots:
        return
    n_slots = len(slots)
    n_needed = len(hasil)

    if n_needed > n_slots:
        last_el = doc.paragraphs[slots[-1]]._element
        for _ in range(n_needed - n_slots):
            new_el = deepcopy(last_el)
            _strip_bookmarks(new_el)
            last_el.addnext(new_el)
            last_el = new_el
    elif n_needed < n_slots:
        for pi in reversed(slots[n_needed:]):
            doc.paragraphs[pi]._element.getparent().remove(doc.paragraphs[pi]._element)

    # Replace each {{HASIL_N}} in order
    idx = 0
    for para in list(doc.paragraphs):
        m = re.search(r"\{\{HASIL_\d+\}\}", para.text)
        if not m:
            continue
        val = hasil[idx] if idx < len(hasil) else ""
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = val
        idx += 1


def _rebuild_table(doc, pelaksana: list):
    """Clone or trim table rows to match pelaksana count."""
    tables = doc.tables
    if not tables:
        return
    table = tables[0]

    existing_data = list(table.rows)[1:]  # skip header
    target = len(pelaksana)

    if len(existing_data) > target:
        for row in reversed(existing_data[target:]):
            table._tbl.remove(row._tr)
    elif len(existing_data) < target:
        template_row = existing_data[-1] if existing_data else table.rows[-1]
        for _ in range(target - len(existing_data)):
            clone = deepcopy(template_row._tr)
            _strip_bookmarks(clone)
            table._tbl.append(clone)

    data_rows = list(table.rows)[1:]
    for i, p in enumerate(pelaksana):
        cells = data_rows[i].cells
        for ci, val in enumerate([f"{i+1}.", p.nama, p.peran_tugas]):
            if ci < len(cells):
                cell = cells[ci]
                for pi, para in enumerate(list(cell.paragraphs)):
                    pPr = para._element.find(qn('w:pPr'))
                    if pPr is not None:
                        numPr = pPr.find(qn('w:numPr'))
                        if numPr is not None:
                            pPr.remove(numPr)
                    para.paragraph_format.space_before = 0
                    para.paragraph_format.space_after = 0
                    para.paragraph_format.line_spacing = 1.0
                    for run in para.runs:
                        run.text = ""
                    if pi > 0 and not para.text.strip():
                        para._element.getparent().remove(para._element)
                if cell.paragraphs and cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].text = val


def _rebuild_ttd(doc, pelaksana: list, nama_ttd: list):
    """Clone signature rows so all names appear under 'Pelaksana Kegiatan,'.

    IMPORTANT: This must run BEFORE _replace_in_doc replaces {{PELAKSANA_N}}
    placeholders, because it detects slots by searching for the literal
    {{PELAKSANA_N}} pattern in paragraph text.
    """
    import re
    used = list(dict.fromkeys(nama_ttd))

    # Find first PELAKSANA_N slot AFTER P[4] (Dari section has them inline)
    first_slot = None
    for pi, para in enumerate(doc.paragraphs):
        if pi <= 4:
            continue
        if re.search(r"\{\{PELAKSANA_\d+\}\}", para.text):
            first_slot = pi
            break
    if first_slot is None:
        return

    # Collect consecutive PELAKSANA_N slots from first_slot onward
    slots = []
    for pi in range(first_slot, len(doc.paragraphs)):
        if re.search(r"\{\{PELAKSANA_\d+\}\}", doc.paragraphs[pi].text):
            slots.append(pi)

    if not slots:
        return
    n_slots = len(slots)
    n_needed = len(used)

    if n_needed > n_slots:
        last_el = doc.paragraphs[slots[-1]]._element
        for _ in range(n_needed - n_slots):
            new_el = deepcopy(last_el)
            _strip_bookmarks(new_el)
            last_el.addnext(new_el)
            last_el = new_el
    elif n_needed < n_slots:
        for pi in reversed(slots[n_needed:]):
            doc.paragraphs[pi]._element.getparent().remove(doc.paragraphs[pi]._element)

    # Clean up empty runs in all TTD paragraphs
    for pi in range(first_slot, len(doc.paragraphs)):
        para = doc.paragraphs[pi]
        para.runs[:] = [r for r in para.runs if r.text.strip() or '<w:drawing>' in r._element.xml]
        para.paragraph_format.line_spacing = 1.0

    # Replace each slot in order (find them dynamically after tree changes)
    current_slots = []
    for pi in range(first_slot, len(doc.paragraphs)):
        if re.search(r"\{\{PELAKSANA_\d+\}\}", doc.paragraphs[pi].text):
            current_slots.append(doc.paragraphs[pi])

    for idx, para in enumerate(current_slots):
        name = used[idx] if idx < len(used) else ""
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = name
        if not name:
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                numPr = pPr.find(qn('w:numPr'))
                if numPr is not None:
                    pPr.remove(numPr)


def _normalize_waktu(raw: str) -> str:
    """Strip descriptive time words (pagi, sore, WIB, dll), keep only HH:MM."""
    import re
    if not raw or raw.strip().lower() == "selesai":
        return raw
    m = re.search(r'(\d{1,2})[.:](\d{2})', raw)
    if m:
        return f"{m.group(1).zfill(2)}:{m.group(2)}"
    return raw


def render_laporan(laporan: LaporanPerdin, template_path: str, output_path: str):
    doc = Document(template_path)

    # ── Step 0: center title ──
    if doc.paragraphs:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Emu
        p0 = doc.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.left_indent = 0
        p0.paragraph_format.right_indent = Emu(1270)
        p0.paragraph_format.first_line_indent = 0

    # ── Step 1: normalise waktu ──
    mulai = _normalize_waktu(laporan.kegiatan_waktu_mulai)
    selesai = _normalize_waktu(laporan.kegiatan_waktu_selesai)

    # ── Step 2: combine waktu ──
    waktu = (
        f"{mulai} WIB - {selesai} WIB"
        if selesai and selesai != "selesai"
        else f"{mulai} WIB"
        if mulai
        else selesai
    )

    # Clean up any leftover waktu suffix in P[19]
    if len(doc.paragraphs) > 19:
        import re
        p19 = doc.paragraphs[19]
        for run in p19.runs:
            run.text = re.sub(r'[^\x00-\x7F]+\s*selesai', '', run.text)

    # ── Step 3: rebuild dynamic sections FIRST (before placeholder replace) ──
    # These functions detect slots by searching for literal {{HASIL_N}} and
    # {{PELAKSANA_N}} patterns. If we replace those placeholders first, the
    # slot detection fails and paragraphs get corrupted.
    _rebuild_hasil(doc, laporan.hasil)
    _rebuild_ttd(doc, laporan.pelaksana, laporan.nama_ttd)

    # ── Step 4: replace simple {{}} placeholders ──
    # NOTE: {{HASIL_N}} and {{PELAKSANA_N}} are already handled by the
    # rebuild functions above. Do NOT include them here.
    replacements = {
        "{{KEPADA}}": laporan.kepada,
        "{{NOMOR_ST}}": laporan.nomor_st,
        "{{TANGGAL_ST}}": laporan.tanggal_st,
        "{{MAKSUD_TUJUAN}}": laporan.maksud_tujuan,
        "{{KEGIATAN_DESKRIPSI}}": laporan.kegiatan_deskripsi,
        "{{KEGIATAN_WAKTU}}": waktu,
        "{{KEGIATAN_TEMPAT}}": laporan.kegiatan_tempat,
        "{{HARI_TANGGAL}}": laporan.hari_tanggal,
        "{{HASIL_INTRO}}": laporan.hasil_intro,
        "{{PENUTUP}}": laporan.penutup,
        "{{TEMPAT_TANGGAL_TTD}}": laporan.tempat_tanggal_ttd,
        "{{TEMBUSAN}}": laporan.tembusan,
    }

    _replace_in_doc(doc, replacements)

    # ── Step 5: rebuild Dari section ──
    _rebuild_dari_section(doc, laporan.pelaksana, laporan.kepada,
                          laporan.tembusan, laporan.hari_tanggal)

    # ── Step 6: compress spacing between sections ──
    for pi in range(28, 77):
        if pi < len(doc.paragraphs) and not doc.paragraphs[pi].text.strip():
            doc.paragraphs[pi].paragraph_format.space_before = 0
            doc.paragraphs[pi].paragraph_format.space_after = 0

    # ── Step 7: rebuild table ──
    _rebuild_table(doc, laporan.pelaksana)

    # ── Step 7b: collapse gap between table and "Kegiatan" heading ──
    for pi in range(13, max(18, len(doc.paragraphs))):
        if pi >= len(doc.paragraphs):
            break
        txt = doc.paragraphs[pi].text.strip()
        if txt == "Kegiatan":
            from docx.shared import Pt
            doc.paragraphs[pi].paragraph_format.space_before = Pt(6)
            doc.paragraphs[pi].paragraph_format.space_after = Pt(2)
        elif pi == 17 and txt:
            doc.paragraphs[pi].paragraph_format.space_before = Pt(2)
            doc.paragraphs[pi].paragraph_format.space_after = 0
        elif not txt:
            doc.paragraphs[pi].paragraph_format.space_before = 0
            doc.paragraphs[pi].paragraph_format.space_after = 0
            for run in doc.paragraphs[pi].runs:
                run.text = ""

    # ── Step 8: save ──
    # NOTE: drawing elements (kop surat/gambar) are NOT removed to keep
    # the document valid. Removing XML drawings while keeping their rels
    # references produces broken OOXML that Word cannot open.
    doc.save(output_path)
